import calendar
import re
from datetime import date
from decimal import Decimal
from io import BytesIO

from docx import Document
from fastapi import HTTPException, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.models import Grant, GrantAward, PSRDueDate, PSRProject

MONTH_LABEL_FMT = "%B %Y"
_MONTH_RE = re.compile(r"(\d{4})-(\d{2})")


def parse_month(month_str: str) -> tuple[int, int]:
    match = _MONTH_RE.fullmatch(month_str or "")
    if not match:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, "month must be in YYYY-MM format")
    year, month = int(match.group(1)), int(match.group(2))
    if not 1 <= month <= 12:
        raise HTTPException(status.HTTP_422_UNPROCESSABLE_ENTITY, "month must be between 01 and 12")
    return year, month


def month_bounds(year: int, month: int) -> tuple[date, date]:
    last_day = calendar.monthrange(year, month)[1]
    return date(year, month, 1), date(year, month, last_day)


def fetch_monthly_report(db: Session, year: int, month: int) -> dict:
    start, end = month_bounds(year, month)

    psr_submitted_rows = db.execute(
        select(PSRDueDate, PSRProject)
        .join(PSRProject, PSRDueDate.project_id == PSRProject.id)
        .where(PSRDueDate.due_date >= start, PSRDueDate.due_date <= end, PSRDueDate.submitted.is_(True))
        .order_by(PSRDueDate.due_date.asc())
    ).all()

    psr_due_rows = db.execute(
        select(PSRDueDate, PSRProject)
        .join(PSRProject, PSRDueDate.project_id == PSRProject.id)
        .where(PSRDueDate.due_date >= start, PSRDueDate.due_date <= end)
        .order_by(PSRDueDate.due_date.asc())
    ).all()

    psr_performance_ending = db.scalars(
        select(PSRProject)
        .where(PSRProject.performance_end_date >= start, PSRProject.performance_end_date <= end)
        .order_by(PSRProject.performance_end_date.asc())
    ).all()

    grants_awarded = db.scalars(
        select(GrantAward)
        .where(GrantAward.award_date >= start, GrantAward.award_date <= end)
        .order_by(GrantAward.award_date.asc())
    ).all()

    grants_expiring = db.scalars(
        select(Grant)
        .where(
            Grant.withdrawn.is_(False),
            Grant.current_exp_date >= start,
            Grant.current_exp_date <= end,
        )
        .order_by(Grant.current_exp_date.asc())
    ).all()

    def _psr_item(due: PSRDueDate, project: PSRProject) -> dict:
        return {
            "due_date_id": due.id,
            "project_id": project.id,
            "project_name": project.project_name,
            "category": project.category,
            "grantor": project.grantor,
            "due_date": due.due_date,
            "submitted": due.submitted,
            "submitted_date": due.submitted_date,
        }

    grants_awarded_amount = sum((a.amount for a in grants_awarded if a.amount is not None), Decimal("0"))

    return {
        "month": f"{year:04d}-{month:02d}",
        "month_label": date(year, month, 1).strftime(MONTH_LABEL_FMT),
        "psr_submitted": [_psr_item(d, p) for d, p in psr_submitted_rows],
        "grants_awarded": list(grants_awarded),
        "psr_due": [_psr_item(d, p) for d, p in psr_due_rows],
        "psr_performance_ending": [
            {
                "project_id": p.id,
                "project_name": p.project_name,
                "category": p.category,
                "grantor": p.grantor,
                "performance_end_date": p.performance_end_date,
            }
            for p in psr_performance_ending
        ],
        "grants_expiring": list(grants_expiring),
        "psr_submitted_count": len(psr_submitted_rows),
        "grants_awarded_count": len(grants_awarded),
        "grants_awarded_amount": grants_awarded_amount,
        "psr_due_count": len(psr_due_rows),
        "psr_performance_ending_count": len(psr_performance_ending),
        "grants_expiring_count": len(grants_expiring),
    }


def build_monthly_report_filename(month: str, ext: str) -> str:
    return f"monthly_report_{month}.{ext}"


def _fmt_date(d: date | None) -> str:
    return d.strftime("%m/%d/%Y") if d else "—"


def _fmt_currency(amount: Decimal | None) -> str:
    return f"${amount:,.2f}" if amount is not None else "—"


def _add_table(document: Document, headers: list[str], rows: list[list[str]], empty_message: str) -> None:
    if not rows:
        paragraph = document.add_paragraph(empty_message)
        paragraph.runs[0].italic = True
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def render_monthly_report_docx(data: dict, generated_by_name: str) -> bytes:
    document = Document()
    document.add_heading("Monthly Report", level=0)
    subheader = document.add_paragraph(
        f"{data['month_label']} — LA County Parks & Recreation Grants Management System"
    )
    subheader.runs[0].italic = True

    document.add_heading("Status Reports Submitted", level=1)
    _add_table(
        document,
        ["Project Name", "Category", "Grantor", "Due Date", "Submitted Date"],
        [
            [i["project_name"], i["category"], i["grantor"] or "—", _fmt_date(i["due_date"]), _fmt_date(i["submitted_date"])]
            for i in data["psr_submitted"]
        ],
        "No status reports submitted this month.",
    )

    document.add_heading("Grants Awarded", level=1)
    _add_table(
        document,
        ["Project Name", "Grantor", "Award Date", "Amount"],
        [
            [a.project_name, a.grantor or "—", _fmt_date(a.award_date), _fmt_currency(a.amount)]
            for a in data["grants_awarded"]
        ],
        "No grants awarded this month.",
    )

    document.add_heading("Status Reports Due", level=1)
    _add_table(
        document,
        ["Project Name", "Category", "Grantor", "Due Date", "Status"],
        [
            [
                i["project_name"],
                i["category"],
                i["grantor"] or "—",
                _fmt_date(i["due_date"]),
                "Submitted" if i["submitted"] else "Not submitted",
            ]
            for i in data["psr_due"]
        ],
        "No status reports due this month.",
    )

    document.add_heading("Status Reports — Performance Period Ending", level=1)
    _add_table(
        document,
        ["Project Name", "Category", "Grantor", "Performance End Date"],
        [
            [i["project_name"], i["category"], i["grantor"] or "—", _fmt_date(i["performance_end_date"])]
            for i in data["psr_performance_ending"]
        ],
        "No status report performance periods end this month.",
    )

    document.add_heading("Grants Expiring", level=1)
    _add_table(
        document,
        ["Project Name", "Grantor", "District", "Expiration Date", "Grant Amount"],
        [
            [
                g.project_name,
                g.grantor or "—",
                str(g.district) if g.district is not None else "—",
                _fmt_date(g.current_exp_date),
                _fmt_currency(g.grant_amount),
            ]
            for g in data["grants_expiring"]
        ],
        "No grants expiring this month.",
    )

    footer = document.add_paragraph(f"Generated on {date.today().strftime('%m/%d/%Y')} by {generated_by_name}.")
    footer.runs[0].italic = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
