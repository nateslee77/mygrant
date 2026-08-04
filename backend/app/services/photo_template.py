import io
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageOps

LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "logo.png"

STATUS_LABELS = {
    "completed": "COMPLETED PROJECT PHOTOS",
    "in_progress": "IN PROGRESS PROJECT PHOTOS",
}

MARGIN = Inches(0.4)

# Photos are center-cropped to this width:height ratio so every photo has a
# predictable, bounded height regardless of its original orientation --
# without this, a single portrait phone photo could blow out to 6+ inches
# tall and push everything onto extra pages.
PHOTO_ASPECT_RATIO = 3 / 2  # width / height

# Target photo width per row-count. Tuned (with the fixed aspect ratio above)
# so 1-3 rows -- i.e. 2 through 6 photos -- always fit on a single landscape
# page alongside the footer block, with a safety margin for line wrapping.
WIDTH_BY_ROWS = {1: Inches(4.6), 2: Inches(4.0), 3: Inches(2.5)}


def _crop_to_aspect(image: Image.Image, target_ratio: float) -> Image.Image:
    w, h = image.size
    current_ratio = w / h
    if current_ratio > target_ratio:
        new_w = round(h * target_ratio)
        left = (w - new_w) // 2
        box = (left, 0, left + new_w, h)
    else:
        new_h = round(w / target_ratio)
        top = (h - new_h) // 2
        box = (0, top, w, top + new_h)
    return image.crop(box)


def _normalize_image(image_bytes: bytes) -> bytes:
    """Re-encode through Pillow into a plain baseline JPEG, cropped to a
    consistent aspect ratio. python-docx's own image-format sniffer is much
    stricter than a real image viewer -- many real-world camera/phone JPEGs
    (progressive encoding, EXIF-only headers, CMYK, etc.) fail it with
    UnrecognizedImageError even though they're perfectly valid images.
    Round-tripping through Pillow also applies EXIF orientation so photos
    come out right-side-up regardless of how the camera stored rotation,
    and the crop keeps every photo's rendered height predictable so the
    whole document reliably fits on one page."""
    image = Image.open(io.BytesIO(image_bytes))
    image = ImageOps.exif_transpose(image)
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    image = _crop_to_aspect(image, PHOTO_ASPECT_RATIO)
    out = io.BytesIO()
    image.save(out, format="JPEG", quality=90)
    return out.getvalue()


def _set_cell_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_pr.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"})
        borders.append(el)
    tbl_pr.append(borders)


def generate_photo_template(
    project_name: str,
    status: str,
    photos: list[tuple[bytes, str]],  # (image_bytes, label)
) -> bytes:
    if not (2 <= len(photos) <= 6):
        raise ValueError("Must have between 2 and 6 photos")

    document = Document()

    # Tighten default paragraph spacing document-wide so caption/title lines
    # don't eat into the page-height budget the sizing math above assumes.
    normal_style = document.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = MARGIN

    rows = (len(photos) + 1) // 2
    photo_width = WIDTH_BY_ROWS.get(rows, Inches(2.5))

    table = document.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_borders_none(table)

    for i, (image_bytes, label) in enumerate(photos):
        row, col = divmod(i, 2)
        cell = table.cell(row, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        img_paragraph = cell.paragraphs[0]
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_paragraph.add_run()
        run.add_picture(io.BytesIO(_normalize_image(image_bytes)), width=photo_width)

        caption_paragraph = cell.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label:
            caption_run = caption_paragraph.add_run(label)
            caption_run.font.size = Pt(10)

    footer_table = document.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    _set_cell_borders_none(footer_table)
    footer_table.columns[0].width = Inches(1.0)
    footer_table.columns[1].width = Inches(8.5)
    footer_table.cell(0, 0).width = Inches(1.0)
    footer_table.cell(0, 1).width = Inches(8.5)

    logo_cell = footer_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(0.7))

    text_cell = footer_table.cell(0, 1)
    subtitle_paragraph = text_cell.paragraphs[0]
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_paragraph.add_run(STATUS_LABELS.get(status, STATUS_LABELS["in_progress"]))
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)

    title_paragraph = text_cell.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(project_name)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def validate_image(image_bytes: bytes) -> None:
    """Raises if the bytes aren't a readable image."""
    Image.open(io.BytesIO(image_bytes)).verify()
