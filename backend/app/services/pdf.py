import re
from datetime import date
from decimal import Decimal
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from app.models.models import Grant, GrantAward, GrantNote
from app.services.grant_status import compute_status

TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "templates"
_env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))

_STATUS_CLASS = {"Active": "active", "Closed": "closed", "Withdrawn": "withdrawn"}


def _safe_filename_part(text: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return slug or "grant"


def render_grant_pdf(grant: Grant, notes: list[GrantNote], status: str, generated_by_name: str) -> bytes:
    from weasyprint import HTML  # imported lazily: requires the GTK3 runtime (Pango/Cairo/GObject),
    # not available in every environment (notably plain Windows without it installed).

    template = _env.get_template("grant_snapshot.html")
    amount_display = f"${grant.grant_amount:,.2f}" if grant.grant_amount is not None else "—"
    sharepoint_is_url = bool(grant.sharepoint_link) and grant.sharepoint_link.strip().lower().startswith(("http://", "https://"))

    html_content = template.render(
        grant=grant,
        notes=notes,
        status=status,
        status_class=_STATUS_CLASS.get(status, "closed"),
        amount_display=amount_display,
        sharepoint_is_url=sharepoint_is_url,
        generated_date=date.today().strftime("%m/%d/%Y"),
        generated_by=generated_by_name,
    )
    return HTML(string=html_content).write_pdf()


def build_snapshot_filename(project_name: str) -> str:
    return f"{_safe_filename_part(project_name)}_snapshot_{date.today().isoformat()}.pdf"


def render_grant_awards_report_pdf(awards: list[GrantAward], generated_by_name: str) -> bytes:
    from weasyprint import HTML  # imported lazily: see note in render_grant_pdf above.

    template = _env.get_template("grant_awards_report.html")
    total_amount = sum((a.amount for a in awards if a.amount is not None), Decimal("0"))

    html_content = template.render(
        awards=awards,
        total_count=len(awards),
        total_amount_display=f"${total_amount:,.2f}",
        generated_date=date.today().strftime("%m/%d/%Y"),
        generated_by=generated_by_name,
    )
    return HTML(string=html_content).write_pdf()


def build_grant_awards_report_filename() -> str:
    return f"grants_awarded_report_{date.today().isoformat()}.pdf"


def _summarize_active_grants(grants: list[Grant]) -> tuple[list[Grant], Decimal, list[dict]]:
    active_rows = [g for g in grants if compute_status(g) == "Active"]
    active_total = sum((g.grant_amount for g in active_rows if g.grant_amount is not None), Decimal("0"))

    by_grantor: dict[str, dict] = {}
    for g in active_rows:
        key = g.grantor or "Unspecified"
        entry = by_grantor.setdefault(key, {"grantor": key, "count": 0, "total": Decimal("0")})
        entry["count"] += 1
        entry["total"] += g.grant_amount or Decimal("0")
    grantor_rows = sorted(by_grantor.values(), key=lambda r: r["total"], reverse=True)

    return active_rows, active_total, grantor_rows


def render_grants_report_pdf(
    grants: list[Grant], report_type: str, generated_by_name: str, filters_description: str | None = None
) -> bytes:
    from weasyprint import HTML  # imported lazily: see note in render_grant_pdf above.

    template = _env.get_template("grants_report.html")
    rows = [(g, compute_status(g)) for g in grants]

    if report_type == "summary":
        active_rows, active_total, grantor_rows = _summarize_active_grants(grants)

        html_content = template.render(
            report_type="summary",
            filters_description=filters_description,
            active_count=len(active_rows),
            active_total_display=f"${active_total:,.2f}",
            grantor_rows=[
                {"grantor": r["grantor"], "count": r["count"], "total_display": f"${r['total']:,.2f}"}
                for r in grantor_rows
            ],
            generated_date=date.today().strftime("%m/%d/%Y"),
            generated_by=generated_by_name,
        )
    else:
        # `rows` arrives pre-ordered by the caller (grants.py preserves the requested
        # grant_ids order, which mirrors whatever sort was applied on the All Grants
        # page, or falls back to alphabetical when no specific grants were requested)
        # -- don't re-sort here.
        total_amount = sum((g.grant_amount for g, _ in rows if g.grant_amount is not None), Decimal("0"))
        html_content = template.render(
            report_type="full",
            filters_description=filters_description,
            rows=rows,
            total_count=len(rows),
            total_amount_display=f"${total_amount:,.2f}",
            generated_date=date.today().strftime("%m/%d/%Y"),
            generated_by=generated_by_name,
        )

    return HTML(string=html_content).write_pdf()


def build_grants_report_filename(report_type: str, ext: str = "pdf") -> str:
    return f"grants_{report_type}_report_{date.today().isoformat()}.{ext}"


def render_grants_summary_report_docx(
    grants: list[Grant], generated_by_name: str, filters_description: str | None = None
) -> bytes:
    from io import BytesIO

    from docx import Document

    active_rows, active_total, grantor_rows = _summarize_active_grants(grants)

    document = Document()
    document.add_heading("All Grants — Summary Report", level=0)
    subheader = document.add_paragraph("LA County Parks & Recreation Grants Management System")
    subheader.runs[0].italic = True
    if filters_description:
        filters_para = document.add_paragraph(f"Filters applied: {filters_description}")
        filters_para.runs[0].italic = True

    stats = document.add_paragraph()
    stats.add_run(f"Active Grants: {len(active_rows)}").bold = True
    stats.add_run("     ")
    stats.add_run(f"Active Grants Under Management: {len(active_rows)} · ${active_total:,.2f}").bold = True

    document.add_heading("Active Grants by Grantor", level=1)
    if grantor_rows:
        table = document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        header_cells = table.rows[0].cells
        for i, header in enumerate(["Grantor", "# of Grants", "Total Amount"]):
            header_cells[i].text = header
            for run in header_cells[i].paragraphs[0].runs:
                run.bold = True
        for r in grantor_rows:
            cells = table.add_row().cells
            cells[0].text = r["grantor"]
            cells[1].text = str(r["count"])
            cells[2].text = f"${r['total']:,.2f}"
    else:
        empty = document.add_paragraph("No active grants recorded.")
        empty.runs[0].italic = True

    footer = document.add_paragraph(f"Generated on {date.today().strftime('%m/%d/%Y')} by {generated_by_name}.")
    footer.runs[0].italic = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_monthly_report_pdf(data: dict, generated_by_name: str) -> bytes:
    from weasyprint import HTML  # imported lazily: see note in render_grant_pdf above.

    template = _env.get_template("monthly_report.html")
    html_content = template.render(
        **data,
        grants_awarded_amount_display=f"${data['grants_awarded_amount']:,.2f}",
        generated_date=date.today().strftime("%m/%d/%Y"),
        generated_by=generated_by_name,
    )
    return HTML(string=html_content).write_pdf()
